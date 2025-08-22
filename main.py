from openai import OpenAI
from docx import Document
import random
import datetime
import re
from dotenv import load_dotenv
import os

# Load environment variables from .env file
load_dotenv()

# Get the API key
api_key = os.getenv("OPENAI_API_KEY")

# ---------- Setup ----------
client = OpenAI(api_key=api_key)

# ---------- Extract Q&A from Benchmark DOCX ----------
def extract_qa_from_table(file_path, heading_text):
    doc = Document(file_path)
    qa_pairs = []
    found_heading = False
    norm_heading = re.sub(r"\s+", " ", heading_text.strip().lower())

    block_iter = iter(doc.element.body.iterchildren())
    for block in block_iter:
        # look for heading paragraph
        if block.tag.endswith("p"):
            para = block.xpath(".//w:t")
            text = "".join([t.text for t in para if t.text])
            norm_text = re.sub(r"\s+", " ", text.strip().lower())
            if norm_heading in norm_text:
                found_heading = True
        # now look for the table
        elif block.tag.endswith("tbl") and found_heading:
            for table in doc.tables:
                if table._element == block:
                    for row in table.rows:
                        if len(row.cells) >= 2:
                            question = row.cells[0].text.strip()
                            answer = row.cells[1].text.strip()
                            if question and answer:
                                qa_pairs.append((question, answer))
                    return qa_pairs
    return qa_pairs

# ---------- Read Transcript Style ----------
def get_transcript_style(file_path):
    doc = Document(file_path)
    all_text = " ".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    return all_text

# ---------- Random Date Generator ----------
def random_date_2025():
    start = datetime.date(2025, 1, 1)
    end = datetime.date.today()
    delta = end - start
    rand_day = random.randrange(delta.days)
    return start + datetime.timedelta(days=rand_day)

# ---------- AI Rewrite ----------
def rewrite_answer(question, benchmark_answer, transcript_style):
    context = (
        "You are rewriting benchmark answers for BSBTWK502 Manage team effectiveness "
         "This is a competency conversation document for BSBTWK502 Manage team effectiveness "
    "on the course CHC52021. You are rewriting benchmark answers provided in Section C. "
    "Rules to follow strictly:\n"
    "1. Give a new rewritten answer that is guided by the benchmark answer, but paraphrased "
    "based on the words and style in the student's call recording transcript.\n"
    "2. The revised answer must keep the benchmark criteria while sounding like the transcript.\n"
    "3. If the benchmark answer is long, the rewritten one must also be long and detailed, "
    "reflecting the transcript style.\n"
    "4. Use the student's name based on the filename of the transcript (e.g., from 'Sample call record transcript.docx').\n"
    "5. If the question asks for 3 examples or 3 occasions, the answer must include exactly 3 examples, "
    "formatted as 'Example 1: ...', 'Example 2: ...', 'Example 3: ...'.\n"
    "6. Keep terminologies with capitalised first letters (like Team Meetings, Key Performance Indicators) "
    "capitalised in the rewritten answer.\n"
    f"7. Insert a random date between past and current in 2025, like {random_date_2025()}, where suitable."
    )

    prompt = (
        f"Transcript Style Sample: {transcript_style}\n\n"
        f"Question: {question}\n\n"
        f"Benchmark Answer: {benchmark_answer}\n\n"
        "Rewrite the Benchmark Answer in the style of the Transcript."
    )

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": context},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content

# ---------- Main ----------
if __name__ == "__main__":
    benchmark_file = "Sample file to Paraphrase.docx"
    transcript_file = "Sample call record transcript.docx"
    heading = "CHCECE042 Foster holistic early childhood learning, development and wellbeing"

    qa_pairs = extract_qa_from_table(benchmark_file, heading)
    transcript_style = get_transcript_style(transcript_file)

    # Create new DOCX with table layout
    new_doc = Document()
    new_doc.add_heading("Rewritten Competency Q&A - CHCECE042", level=1)

    table = new_doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Question"
    hdr_cells[1].text = "Rewritten Answer"

    for q, benchmark_ans in qa_pairs:
        if "question" in q.lower():   # ✅ check if 'Question' is in the string
            rewritten = rewrite_answer(q, benchmark_ans, transcript_style)
            print('rewrite answer')
            print(rewritten)

            row_cells = table.add_row().cells
            row_cells[0].text = q
            row_cells[1].text = rewritten

    output_file = "Rewritten_Benchmark_Answers.docx"
    new_doc.save(output_file)
    print(f"✅ New DOCX file created: {output_file}")
